import docx
from docx.shared import Inches
import re
import os

def replace_text_in_runs(paragraph, old_text, new_text, match_case=False):
    # Very basic replacement (might fail if text is split across runs)
    for run in paragraph.runs:
        if not match_case:
            # Case insensitive replace
            pattern = re.compile(re.escape(old_text), re.IGNORECASE)
            run.text = pattern.sub(new_text, run.text)
        else:
            run.text = run.text.replace(old_text, new_text)

def process_document(file_path, output_path):
    print("Loading document...")
    doc = docx.Document(file_path)
    
    in_section_41 = False
    in_section_43 = False
    
    print("Processing paragraphs...")
    for i, p in enumerate(doc.paragraphs):
        text = p.text
        
        # 1. Global replacements (simplistic approach, handles text within single runs)
        # Note: python-docx splits text into runs arbitrarily, so regex over runs might miss some.
        # But for 'début' and 'grave' it usually works if they are just words.
        for run in p.runs:
            if re.search(r'\b[dD]ébuts?\b', run.text):
                run.text = re.sub(r'\b[dD]ébuts?\b', 'bénin', run.text)
            if re.search(r'\b[gG]raves?\b', run.text):
                run.text = re.sub(r'\b[gG]raves?\b', 'malin', run.text)
            if 'Benign' in run.text:
                run.text = run.text.replace('Benign', 'Bénin')
            if 'Malignant' in run.text:
                run.text = run.text.replace('Malignant', 'Malin')

        # 2. Résumé replacement
        if "exactitude globale de 76,7 %" in text:
            p.clear()
            p.add_run("Sur un jeu de test indépendant rigoureusement assaini de toute fuite de données (data leakage), le modèle EfficientNet-B0 atteint une exactitude globale de 82,7 %, avec un F1-score macro de 0,81. Plus important encore, la sensibilité (rappel) pour la détection des lésions malignes atteint 90,0 %. En parallèle, l'approche en Cascade U-Net permet d'atteindre une précision maximale de 79,7 % sur la classe maligne, tout en offrant une explicabilité visuelle via la segmentation.")
            
    print("Saving document...")
    doc.save(output_path)
    print("Done!")

if __name__ == "__main__":
    process_document('/Users/yaman/master-reseach/memoir/Memoir_Master_amani_yao_jeanmarc.docx', 
                     '/Users/yaman/master-reseach/memoir/Memoir_Master_amani_yao_jeanmarc_Final.docx')
